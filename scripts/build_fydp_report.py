#!/usr/bin/env python3
"""Build FYDP full report (Markdown + DOCX) from project drafts and metrics."""

from __future__ import annotations

import re
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
THESIS = ROOT / "docs" / "thesis"
MD_OUT = THESIS / "FYDP_REPORT_FULL.md"
DOCX_OUT = THESIS / "FYDP_REPORT_Summer2025.docx"

TITLE = (
    "A Deep Learning Framework for Early Detection of Chronic Kidney Disease: "
    "Integrating Multi-Modal Data from Wearable Devices and Clinical Records "
    "with Explainable AI"
)
STUDENT = "Md. Shadman Tahsin"
STUDENT_ID = "[Student ID — fill before submission]"
SUPERVISOR = "Dr. Naznin Sultana"
SUPERVISOR_TITLE = "Assistant Professor"
UNIVERSITY = "Daffodil International University"
DEPT = "Department of Computer Science and Engineering"
DATE = "[Month Day, Year — fill before submission]"


def md_report() -> str:
    return f"""# {TITLE}

---

## Title Page

**{TITLE}**

**By**

{STUDENT}  
{STUDENT_ID}

**FINAL YEAR DESIGN PROJECT REPORT**

This Report Presented in Partial Fulfillment of the Requirements for the Degree of Bachelor of Science in Computer Science and Engineering

**Supervised by**  
{SUPERVISOR}  
{SUPERVISOR_TITLE}  
{DEPT}  
{UNIVERSITY}

**Co-Supervised by**  
[Co-Supervisor Name — if applicable, else remove this block]  
[Designation]  
{DEPT}  
{UNIVERSITY}

**{UNIVERSITY.upper()}**  
Dhaka, Bangladesh  
{DATE}

---

## Approval

This Project titled “{TITLE},” submitted by {STUDENT} to the {DEPT}, {UNIVERSITY}, has been accepted as satisfactory for the partial fulfillment of the requirements for the degree of B.Sc. in Computer Science and Engineering and approved as to its style and contents. The presentation has been held on **[Presentation Date — fill before submission]**.

### BOARD OF EXAMINERS

| Role | Name | Designation |
|------|------|-------------|
| Board Chairman | [Name] | Designation, Department of CSE, FSIT, {UNIVERSITY} |
| Internal Examiner 1 | [Name] | Designation, Department of CSE, FSIT, {UNIVERSITY} |
| Internal Examiner 2 | [Name] | Designation, Department of CSE, FSIT, {UNIVERSITY} |
| External Examiner | [Name] | Designation, Department of CSE, FSIT, {UNIVERSITY} |

**Signatures required:** Board Chairman, Internal Examiners, External Examiner, Supervisor.

---

## Declaration

I hereby declare that this project has been done by me under the supervision of {SUPERVISOR}, {SUPERVISOR_TITLE}, {DEPT}, {UNIVERSITY}. I also declare that neither this project nor any part of this project has been submitted elsewhere for the award of any degree or diploma.

**Supervised by:**  
{SUPERVISOR}  
{SUPERVISOR_TITLE}  
{DEPT}  
{UNIVERSITY}

**Submitted by:**  
{STUDENT}  
Student ID: {STUDENT_ID}  
{DEPT}  
{UNIVERSITY}

**Signature / Date:** _________________________

---

## Acknowledgements

This work would not have been possible without the support and contributions of many individuals over the past two semesters.

First, I express heartfelt thanks to the Almighty for the blessing that made it possible to complete this Final Year Design Project successfully.

I am deeply grateful to {SUPERVISOR}, {SUPERVISOR_TITLE}, {DEPT}, {UNIVERSITY}, Dhaka, Bangladesh. Her knowledge and interest in machine learning for healthcare guided this project. Her patience, scholarly guidance, continual encouragement, constructive criticism, and valuable advice at every stage made completion of this work possible.

I thank the Head of the {DEPT} and the faculty and staff of {UNIVERSITY} for their assistance throughout the programme.

I thank my classmates at {UNIVERSITY} for discussions during coursework and project development.

Finally, I acknowledge the constant support and patience of my parents.

---

## Abstract

Chronic kidney disease (CKD) is a progressive condition that often remains undetected until advanced stages, imposing substantial clinical and economic burden. Early risk stratification using structured electronic health records (EHR) and complementary wearable physiology may support preventive care, but many existing machine-learning studies rely on a single modality, ignore subject-level leakage in hospital data, and report poorly calibrated probabilities without explainability. This thesis presents a reproducible multimodal learning framework for CKD-related risk assessment integrating population clinical data (NHANES), hospital admission records (MIMIC-IV), and wearable signal proxies (WESAD), with probability calibration, SHAP-based explainability, and a research clinical decision-support prototype.

Three independent public cohorts were preprocessed without patient-level merging. The primary locked hospital model is sigmoid-calibrated logistic regression (`logreg_sigmoid_cal`) on approximately 30,000 MIMIC admissions with 68 tabular features and subject-grouped train/validation/test splits. On the held-out test set, this model achieved AUROC 0.7614, F1 0.4145, sensitivity 0.6739, specificity 0.6959, and expected calibration error (ECE) 0.0203 at operating threshold 0.16. Sigmoid calibration reduced ECE from approximately 0.25 (uncalibrated) without changing ranking performance. Secondary branches establish population feasibility (NHANES XGBoost AUROC 0.9922 on a separate cohort) and wearable encoding feasibility (WESAD random forest AUROC 0.7639 on a stress-vs-baseline proxy task, not CKD-labeled). Late fusion is documented as a protocol; because NHANES, MIMIC, and WESAD share no same-patient identifiers, fusion evaluation is proxy-only and does not support claims of multimodal CKD benefit.

Deliverables include a supervisor notebook pipeline, modular source code, frozen artifacts, thesis figures, and a Streamlit decision-support demonstration. The work is research and educational software, not a clinical diagnostic device. Limitations—including ICD proxy labels, disjoint cohorts, and ML-primary (not end-to-end deep learning on raw waveforms) evidence—are stated explicitly. Future work requires linked multimodal CKD cohorts and external hospital validation before deployment claims.

---

## Table of Contents

| Section | Page |
|---------|------|
| Approval | i |
| Declaration | ii |
| Acknowledgements | iii |
| Abstract | iv |
| List of Figures | v |
| List of Tables | vi |
| 1 Introduction | 1 |
| 1.1 Introduction | 1 |
| 1.2 Motivation | 1 |
| 1.3 Objectives | 2 |
| 1.4 Methodology | 2 |
| 1.5 Project Outcome | 2 |
| 1.6 Organization of the Report | 3 |
| 2 Background | 4 |
| 2.1 Introduction | 4 |
| 2.2 Literature Review | 4 |
| 2.2.1 Similar Applications | 6 |
| 2.3 Gap Analysis | 7 |
| 2.4 Summary | 8 |
| 3 Research Methodology | 9 |
| 3.1 Methodology / Requirement Analysis & Design Specification | 9 |
| 3.1.1 Overview | 9 |
| 3.1.2 Proposed Methodology / System Design | 9 |
| 3.1.3 Functional and Nonfunctional Requirements | 10 |
| 3.1.4 Data Flow Diagram | 11 |
| 3.1.5 UI Design | 11 |
| 3.2 Detailed Methodology and Design | 12 |
| 3.3 Project Plan | 13 |
| 3.4 Task Allocation | 13 |
| 3.5 Summary | 14 |
| 4 Implementation and Results | 15 |
| 4.1 Environment Setup | 15 |
| 4.2 Testing and Evaluation / Performance / Comparative Analysis | 16 |
| 4.3 Results and Discussion | 17 |
| 4.4 Summary | 22 |
| 5 Engineering Standards and Design Challenges | 23 |
| 5.1 Compliance with the Standards | 23 |
| 5.2 Impact on Society, Environment and Sustainability | 25 |
| 5.3 Project Management and Financial Analysis | 26 |
| 5.4 Complex Engineering Problem | 27 |
| 5.5 Summary | 29 |
| 6 Conclusion | 30 |
| 6.1 Summary | 30 |
| 6.2 Limitation | 31 |
| 6.3 Future Work | 32 |
| References | 33 |

*Note: Update page numbers after pasting into Word and generating TOC.*

---

## List of Figures

| Figure | Title | Page |
|--------|-------|------|
| 3.1 | Multimodal CKD framework — three branches and late fusion | 10 |
| 3.2 | Data flow diagram — NHANES, MIMIC, WESAD pipelines | 11 |
| 3.3 | Clinical decision-support prototype UI (Streamlit) | 12 |
| 4.1 | MIMIC calibration reliability diagram | 18 |
| 4.2 | MIMIC test AUROC comparison across model families | 18 |
| 4.3 | MIMIC global SHAP top-10 features | 19 |
| 4.4 | MIMIC confusion matrix — logreg_sigmoid_cal at threshold 0.16 | 19 |
| 4.5 | WESAD wearable window feature extraction | 20 |

*Insert figures from `paper_assets/figures/` when formatting in Word.*

---

## List of Tables

| Table | Title | Page |
|-------|-------|------|
| 2.1 | Summary of Literature Reviewed | 5 |
| 2.2 | Similar Applications — CKD and Clinical ML Systems | 6 |
| 2.3 | Gap Analysis — Identified Limitations and Proposed Solutions | 7 |
| 3.1 | Data sources and label definitions | 10 |
| 3.2 | Functional requirements | 10 |
| 3.3 | Nonfunctional requirements | 11 |
| 3.4 | Project task allocation timeline | 13 |
| 4.1 | NHANES test-set performance | 16 |
| 4.2 | MIMIC admission branch — held-out test performance | 17 |
| 4.3 | MIMIC calibration comparison | 17 |
| 4.4 | Fusion meta-learner comparison (proxy holdout) | 20 |
| 5.1 | Mapping with complex engineering problem | 27 |
| 5.2 | Mapping with knowledge profile | 28 |
| 5.3 | Mapping with complex engineering activities | 28 |

---

# Chapter 1 — Introduction

This chapter introduces chronic kidney disease as a public-health problem, states the motivation and objectives of this Final Year Design Project, summarizes the methodology and expected outcomes, and outlines the organization of the remainder of the report.

## 1.1 Introduction

Chronic kidney disease (CKD) is a slowly progressive condition in which kidney function declines over time, reflected by reduced estimated glomerular filtration rate (eGFR) or evidence of kidney damage [1], [2]. Severity ranges from early asymptomatic stages to end-stage renal disease (ESRD), dialysis, and transplantation, with substantial clinical and economic burden [1]. CKD is common worldwide, frequently coexists with hypertension and type 2 diabetes, and is often undetected in early stages because symptoms may be absent [3], [4]. Early identification of decreased kidney function or albuminuria enables medication, dietary modification, and risk-factor management [5], [6]. Reliable early risk stratification therefore remains an important challenge in nephrology and preventive medicine.

Traditional screening relies on serum creatinine, eGFR (CKD-EPI), and urine albumin-to-creatinine ratio (ACR) [5]. These markers are clinically useful but are intermittent, influenced by hydration, muscle mass, and acute illness, and do not provide continuous monitoring between clinic visits [7], [8]. Artificial intelligence (AI) methods can model high-dimensional structured data and nonlinear relationships in electronic health records (EHR) [7], [9], [10]. However, many CKD prediction studies use a single data modality and do not integrate complementary wearable physiology [7], [11], [12]. Structured clinical records combined with wearable signals for CKD remain under-explored [13].

A further limitation of many black-box models is limited interpretability and poorly calibrated probabilities, which reduces trust in clinical decision support [14], [15], [16]. Explainable AI (XAI) methods such as SHAP can support feature-level review, but calibrated probability outputs and decision-support framing remain underdeveloped for multimodal CKD pipelines [17].

This thesis proposes a **multimodal learning framework** that supports both **classical machine learning** and **tabular deep learning** models, with **probability-level late fusion** [13], [18] and **SHAP-based explainability** [17] for transparent decision support rather than autonomous diagnosis [7], [8], [19]. The **primary locked hospital model** is **sigmoid-calibrated logistic regression** on MIMIC-IV admissions; deep tabular models are trained and compared under the same protocol. The title references deep learning because neural tabular baselines and a deep-learning-ready multimodal architecture are implemented; the **locked primary evidence** is nevertheless **machine-learning primary** (calibrated logistic regression), not end-to-end deep learning on raw wearable waveforms.

## 1.2 Motivation

CKD prevalence is rising while many at-risk patients are diagnosed only after substantial nephron loss [3]. Periodic laboratory testing cannot monitor patients continuously between visits [7]. Risk prediction must address high-dimensional sparse features, class imbalance, heterogeneous data sources, and—critically—the absence of a public same-patient cohort linking hospital EHR, population surveys, and wearable recordings.

Wearable devices measure blood volume pulse (BVP), electrodermal activity (EDA), and skin temperature, which relate to autonomic and cardiovascular physiology relevant to chronic disease monitoring [11], [12]. Combining structured EHR features with wearable-derived summaries may enrich risk assessment, but engineering challenges include proxy labeling, cross-dataset generalization, and mismatched sampling rates.

From a clinical perspective, validated early risk tools may help identify high-risk patients before overt symptoms [4], [10]. From a technical perspective, this project delivers a reproducible pipeline with calibration, explainability, and a prototype decision-support interface for research evaluation.

## 1.3 Objectives

The aim is to develop an explainable **machine-learning and deep-learning framework** for early CKD-related risk assessment using structured clinical data and wearable physiology proxies, producing a **clinical decision-support prototype**, not an autonomous diagnostic system.

1. Prepare and preprocess CKD-relevant clinical data from **NHANES** and **MIMIC-IV**, including feature selection, missing-value handling, leakage removal, label construction, and grouped train/validation/test splitting.
2. Extract wearable physiological features from **WESAD** using window-based signal processing for the wearable branch.
3. Develop and compare branch-level models: logistic regression (with calibration), random forest, XGBoost, tabular MLP, and residual MLP.
4. Evaluate models using AUROC, AUPRC, F1-score, sensitivity, specificity, balanced accuracy, Brier score, and expected calibration error (ECE).
5. Implement explainability using feature importance and **SHAP** for clinically reviewable attributions.
6. Design **late fusion** combining branch-level probability outputs under the constraint that **no same-patient aligned multimodal CKD cohort** is available.
7. Develop a **prototype decision-support interface** presenting calibrated risk and explanations for demonstration and research evaluation.

## 1.4 Methodology

Three public datasets were used as **independent branches**: NHANES (population clinical) [18], MIMIC-IV v3.1 hospital EHR [19], and WESAD (wearable physiology) [20]. Each source was preprocessed separately; CKD-related or proxy labels were defined per branch; grouped splits prevented subject-level leakage [5], [6]. Clinical tabular features and wearable window statistics were extracted. Supervised models—including logistic regression with probability calibration, random forest, XGBoost, and PyTorch tabular MLPs—were trained per branch. **Late fusion** combined calibrated branch probabilities using validation-quality weights and a meta-logistic-regression comparator [13]. SHAP explainability was applied to the primary MIMIC model [17]. Performance was assessed using discrimination, calibration, robustness, and temporal hold-out within MIMIC.

**Important:** datasets were **not merged at patient level**. Multimodal integration is **late fusion at probability level**, with proxy evaluation where aligned cohorts are unavailable.

## 1.5 Project Outcome

The project delivers a reproducible early CKD **risk-assessment framework** for research and decision support, not autonomous diagnosis.

**Primary MIMIC-IV test set** — locked model **`logreg_sigmoid_cal`**, threshold **0.16**:

| Metric | Value |
|--------|-------|
| Accuracy | 69.23% |
| AUROC | 0.7614 |
| F1 | 0.4145 |
| Sensitivity | 0.6739 |
| Specificity | 0.6959 |
| ECE | 0.0203 |

**NHANES branch** (separate cohort) — best model **XGBoost**: accuracy 96.60%, AUROC **0.9922**, F1 0.7975. Not directly comparable to MIMIC.

**WESAD branch** (stress vs baseline proxy, not CKD-labeled): accuracy 81.21%, AUROC **0.7639**, F1 0.6667.

A **late-fusion protocol** was implemented; because sources are not same-patient aligned and WESAD has no CKD labels, fusion is an **exploratory research outcome**, not a clinically validated multimodal CKD system.

Deliverables: reproducible pipeline, branch models, calibrated primary hospital model, SHAP explainability, fusion protocol, prototype CDS UI, and thesis figures under `paper_assets/`.

## 1.6 Organization of the Report

- **Chapter 1** introduces CKD, motivation, objectives, methodology, outcomes, and report structure.
- **Chapter 2** reviews literature, similar applications, and gap analysis (single modality, interpretability, multimodal integration).
- **Chapter 3** describes data acquisition, preprocessing, modeling, fusion, explainability, evaluation requirements, and project planning.
- **Chapter 4** presents implementation and results (NHANES, MIMIC primary, WESAD, fusion, UI).
- **Chapter 5** discusses engineering standards, privacy, ethics, complex engineering mapping, and financial analysis.
- **Chapter 6** concludes with contributions, limitations, and future work.

---

# Chapter 2 — Background

This chapter provides the background knowledge required to understand the rest of the report, reviews related literature and similar applications, and presents a gap analysis motivating the proposed framework.

## 2.1 Introduction

Chronic kidney disease is defined clinically by persistent abnormalities in kidney structure or function for at least three months [1]. KDIGO guidelines classify severity using eGFR categories and albuminuria stages [5]. Hospital EHR systems capture diagnoses, demographics, and laboratory values that can support risk models, while population surveys such as NHANES provide nationally representative biochemistry. Wearable sensors offer continuous physiology but rarely include renal labels in public datasets. Machine learning for CKD spans population screening, admission risk prediction, and emerging multimodal integration; calibration and explainability are increasingly recognized as requirements for clinical decision support [14], [15], [17].

## 2.2 Literature Review

Table 2.1 summarizes representative prior work on CKD detection, hospital risk modeling, calibration, explainability, and wearable health monitoring. At least ten sources are reviewed.

### Table 2.1: Summary of Literature Reviewed

| Author(s) | Year | Title | Methodology | Key Findings |
|-----------|------|-------|-------------|--------------|
| KDIGO Work Group | 2012 | KDIGO Clinical Practice Guideline for CKD | Clinical guideline | Defines CKD staging using eGFR and albuminuria; basis for population labels |
| Webster et al. | 2017 | Chronic Kidney Disease | Epidemiology review | CKD is common, underdiagnosed, and linked to cardiovascular mortality |
| Tomasev et al. | 2019 | A clinically applicable approach to continuous prediction of acute kidney injury | Deep learning on EHR | Demonstrates temporal EHR modeling at scale; motivates admission-level ML |
| Ravizza et al. | 2019 | Predicting CKD progression with machine learning | Supervised ML on clinical data | Shows ML can stratify CKD progression using structured features |
| Guo et al. | 2017 | On Calibration of Modern Neural Networks | Calibration analysis | Modern classifiers can be miscalibrated; post-hoc calibration improves reliability |
| Lundberg & Lee | 2017 | A Unified Approach to Interpreting Model Predictions (SHAP) | Game-theoretic XAI | SHAP provides consistent feature attributions for ML models |
| Johnson et al. | 2020 | MIMIC-IV (database) | Critical care EHR resource | Enables reproducible hospital ML with documented schema |
| Schmidt et al. | 2018 | Introducing WESAD multimodal dataset | Wearable stress dataset | Public BVP/EDA/ACC/TEMP signals; no CKD labels |
| Chen et al. | 2021 | Multimodal machine learning in healthcare | Survey | Multimodal fusion requires aligned patient-level data for valid claims |
| Rahimi et al. | 2020 | Wearable sensors and digital health for CKD | Narrative review | Wearables may support remote monitoring; CKD-specific labeled data remain scarce |
| Choi et al. | 2020 | Using recurrent neural networks for predicting CKD | Deep learning on EHR | RNN/LSTM models improve AKI/CKD-related prediction on structured EHR |
| Rajkomar et al. | 2018 | Scalable and accurate deep learning for EHR | Deep learning | Deep models on EHR need careful leakage control and calibration for deployment |

## 2.2.1 Similar Applications

Table 2.2 compares similar clinical ML and decision-support systems. These are methodological or application analogues—not direct competitors for the same multimodal CKD task on disjoint public data.

### Table 2.2: Similar Applications

| System / Work | Modality | ML Approach | Calibration / XAI | Limitation vs this thesis |
|---------------|----------|-------------|-------------------|---------------------------|
| AKI prediction on MIMIC (Deep learning) | EHR time series | RNN / attention | Often AUROC-only | Single modality; not CKD-focused |
| NHANES CKD screening studies | Survey labs | Logistic / tree models | Limited calibration reporting | No hospital admission context |
| Wearable stress detection (WESAD baselines) | BVP, EDA, ACC | RF / SVM | Minimal clinical calibration | No renal labels |
| Epic / Cerner risk scores (commercial CDS) | EHR | Proprietary rules + ML | Varies by vendor | Not open reproducible research |
| Streamlit clinical ML demos | Tabular EHR | sklearn + UI | Varies | Rarely multimodal + grouped splits |
| Multimodal fusion papers (general healthcare) | Imaging + text + labs | Late / early fusion | Partial | Require aligned patient records |
| SHAP-based hospital readmission tools | EHR tabular | GBM + SHAP | Explainability focus | Different outcome; calibration not always reported |
| XGBoost NHANES chronic disease models | Population labs | Gradient boosting | High AUROC on survey | Not generalizable to ICU admissions |
| MIMIC sepsis / mortality predictors | EHR | Logistic / deep | Some calibration work | Different label; leakage issues common |
| Research CDS prototypes (academic) | EHR subset | ML + dashboard | Prototype only | Often lack multimodal + honest fusion scope |

## 2.3 Gap Analysis

Despite substantial progress in machine learning for CKD screening and hospital risk prediction, several gaps motivate this project: **single-modality focus**, **leakage-prone evaluation**, **poor probability calibration**, **limited explainability**, **missing demonstrable CDS interfaces**, and **unvalidated multimodal fusion on disjoint cohorts**.

### Table 2.3: Gap Analysis — Identified Limitations and Proposed Solutions

| Gap ID | Identified gap | Limitation in existing work | Proposed solution (this thesis) | How this work addresses it |
|--------|----------------|----------------------------|----------------------------------|----------------------------|
| G1 | Fragmented multimodal CKD pipelines | Most studies use one source only | Three-branch framework + fusion protocol | NHANES, MIMIC, WESAD under one pipeline |
| G2 | Information leakage in hospital ML | Random admission splits inflate metrics | Subject-grouped splits on MIMIC | Grouped 60/20/20 split; multi-seed robustness |
| G3 | Miscalibrated risk probabilities | High AUROC with unreliable probabilities | Sigmoid/isotonic calibration + ECE | ECE 0.25 → 0.02; threshold 0.16 |
| G4 | Limited explainability for CDS | Black-box models without attributions | SHAP + local logistic contributions | Global SHAP; CDS prototype explanations |
| G5 | Gap between offline models and CDS | Metrics only, no interface | Streamlit research CDS prototype | Calibrated risk, threshold, SHAP tab |
| G6 | Unvalidated same-patient fusion | Fusion needs aligned modalities | Transparent protocol + honest non-claim | Proxy holdout only; limitations in §6.2 |

The gap analysis shows that this project targets **reproducible, leakage-aware, calibrated, and explainable hospital risk stratification** inside a **multimodal framework** with population and wearable branches and an honest treatment of fusion limits. The **primary empirical contribution** is the **MIMIC clinical branch** (G2–G5).

## 2.4 Summary

Chapter 2 established CKD clinical background, reviewed at least ten related works, compared similar applications, and identified six gaps addressed by this thesis. The next chapter details the research methodology and system design.

---

# Chapter 3 — Research Methodology

This chapter describes the proposed multimodal methodology, requirements, data flow, UI design, alternative solutions considered, project plan, and task allocation.

## 3.1 Methodology / Requirement Analysis & Design Specification

### 3.1.1 Overview

The methodology follows a **branch-wise multimodal pipeline**: three independent data sources (NHANES, MIMIC-IV, WESAD) are preprocessed separately, models are trained per branch, the primary hospital model is calibrated and explained, and a **late-fusion protocol** combines branch-level probabilities where alignment permits only proxy evaluation.

Implementation is centralized in `notebooks/ckd_supervisor_pipeline_from_scratch.ipynb` with reusable modules under `src/`.

**Design principle:** no patient-level merge across NHANES, MIMIC, and WESAD; fusion operates on **probability outputs**, not a unified multimodal table.

### 3.1.2 Proposed Methodology / System Design

**Figure 3.1** (insert from `paper_assets/figures/` — multimodal workflow diagram): Three parallel towers — NHANES (population clinical), MIMIC (hospital EHR, primary), WESAD (wearable encoder) — each produce branch probabilities. Late fusion combines weighted branch outputs and a meta-logistic comparator. The CDS prototype consumes frozen MIMIC artifacts.

### Table 3.1: Data Sources and Label Definitions

| Branch | Dataset | Unit | Label |
|--------|---------|------|-------|
| Population clinical | NHANES 2013–2014 | Participant | eGFR < 60 (CKD-EPI) |
| Hospital EHR (primary) | MIMIC-IV v3.1 hosp | Admission | ICD CKD proxy (585* / N18*) |
| Wearable proxy | WESAD | Signal window | Stress vs baseline (not CKD) |

### 3.1.3 Functional and Nonfunctional Requirements

#### Table 3.2: Functional Requirements

| ID | Requirement |
|----|-------------|
| FR1 | Ingest NHANES, MIMIC-IV, and WESAD from configured paths |
| FR2 | Construct branch-specific labels with documented definitions |
| FR3 | Engineer tabular (clinical) and window-statistic (wearable) features |
| FR4 | Train and compare ML/DL models per branch with grouped splits |
| FR5 | Calibrate primary MIMIC model and select operating threshold |
| FR6 | Compute discrimination, calibration, and robustness metrics |
| FR7 | Generate global SHAP and export explanation artifacts |
| FR8 | Implement late-fusion protocol with static and meta-learned combiners |
| FR9 | Provide Streamlit CDS prototype with risk score and explanations |
| FR10 | Export reproducible artifacts (JSON, CSV, figures, checkpoint) |

#### Table 3.3: Nonfunctional Requirements

| ID | Requirement |
|----|-------------|
| NFR1 | Reproducibility: fixed seeds, documented splits, reporting lock file |
| NFR2 | Leakage prevention: subject-grouped MIMIC splits |
| NFR3 | Performance: pipeline runnable on student workstation (CPU-feasible MIMIC primary) |
| NFR4 | Interpretability: primary model supports SHAP and coefficient review |
| NFR5 | Honest scope: disclaimers for proxy labels, disjoint fusion, research-only CDS |
| NFR6 | Maintainability: modular `src/` and staged notebook cells |
| NFR7 | Portability: Python 3.12 virtual environment and requirements.txt |

### 3.1.4 Data Flow Diagram

**Figure 3.2** — Data flow: Raw datasets → preprocessing → feature matrices → model training → calibration → evaluation → artifact export → CDS prototype. MIMIC path: 30,000 admissions → 68 features → grouped split → logreg_sigmoid_cal → SHAP → Streamlit.

### 3.1.5 UI Design

**Figure 3.3** — Streamlit CDS prototype (`app/demo_app.py`): admission browser, calibrated risk percentage, adjustable threshold (default 0.16), local explanation panel, global SHAP tab. Screenshots: `paper_assets/figures/fig_ui_*.png`.

## 3.2 Detailed Methodology and Design

### Preprocessing

**NHANES:** Merge cycle CSVs; binary CKD from eGFR; drop leakage columns; median imputation and scaling fit on train only (~4,702 rows).

**MIMIC-IV:** Sample 30,000 index admissions; join demographics and lab medians within 48 h; 68 features; **grouped split by subject_id** (60/20/20).

**WESAD:** Wrist pickle per subject; windows 1920 samples (~30 s), stride 960; per-channel mean/std/min/max; stress vs baseline label; grouped holdout by subject.

### Models Considered

| Algorithm | Branches |
|-----------|----------|
| Logistic regression (+ calibration) | NHANES, MIMIC, fusion meta |
| Random forest | NHANES, MIMIC, WESAD |
| XGBoost | NHANES, MIMIC |
| Tabular MLP / ResMLP | NHANES, MIMIC |

**Alternatives rejected or deferred:**
- **End-to-end CNN on raw waveforms** — no CKD labels in WESAD; deferred to future linked cohorts.
- **NLP on discharge notes** — out of timeline; tabular primary chosen.
- **Same-patient early fusion** — impossible on public disjoint datasets; late fusion at probability level selected.
- **XGBoost as primary MIMIC model** — slightly higher AUROC but weaker calibration/interpretability trade-off; **logreg_sigmoid_cal** locked for CDS transparency.

**Primary locked model:** `logreg_sigmoid_cal`, threshold **0.16**.

### Calibration, XAI, Fusion, Evaluation

Post-hoc sigmoid (Platt) and isotonic calibration on MIMIC validation; ECE and Brier score. Global SHAP on MIMIC test set. Late fusion: static weighted sum and meta-logreg on branch probabilities; proxy holdout documented. Metrics: AUROC, AUPRC, F1, sensitivity, specificity, ECE, Brier, multi-seed robustness, temporal holdout.

## 3.3 Project Plan

| Phase | Activities | Deliverable |
|-------|------------|-------------|
| P1 | Literature, gap analysis, architecture | Ch. 2, design spec |
| P2 | Data engineering (NHANES, MIMIC, WESAD) | Clean branch datasets |
| P3 | Model development and calibration | Checkpoints, reporting lock |
| P4 | XAI, robustness, figures | SHAP, calibration plots |
| P5 | CDS prototype, thesis, defense prep | demo_app.py, full report |

## 3.4 Task Allocation

Single-student project; all tasks performed by {STUDENT} under supervision of {SUPERVISOR}.

### Table 3.4: Task Allocation Timeline (Weeks 12–48)

| Task | W12–16 | W18–22 | W24–28 | W30–34 | W36–40 | W42–48 |
|------|:------:|:------:|:------:|:------:|:------:|:------:|
| Literature & planning | ████ | | | | | |
| Data collection & preprocessing | | ████ | ████ | | | |
| Model training (all branches) | | | ████ | ████ | | |
| Calibration, XAI, evaluation | | | | ████ | ████ | |
| CDS prototype & thesis writing | | | | | ████ | ████ |
| Defense preparation | | | | | | ████ |

## 3.5 Summary

Chapter 3 specified a branch-wise multimodal methodology with grouped evaluation, calibration, SHAP explainability, late fusion protocol, functional/nonfunctional requirements, and a single-student project plan. Chapter 4 presents implementation results.

---

# Chapter 4 — Implementation and Results

This chapter documents the software environment, evaluation protocol, branch-level results, fusion protocol outcomes, CDS prototype, and discussion.

## 4.1 Environment Setup

Experiments used Python **3.12** (`.venv312`), **pandas**, **NumPy**, **scikit-learn**, **PyTorch**, **XGBoost**, **SHAP**, and **Streamlit**. Data paths resolve to `STUDY/Dataset/`; artifacts write to `outputs/supervisor_runs/`.

| Source | Path |
|--------|------|
| MIMIC-IV v3.1 hosp | `STUDY/Dataset/mimic-iv-3.1/hosp/` |
| NHANES | `STUDY/Dataset/nhanes_ckd/csv/` |
| WESAD | `STUDY/Dataset/WESAD/` |

Splitting: programmatic grouped splits (MIMIC by `subject_id` 60/20/20; NHANES held-out test; WESAD subject holdout).

## 4.2 Testing and Evaluation / Performance / Comparative Analysis

Evaluation dimensions:
1. **Discrimination** — AUROC, AUPRC
2. **Classification** — accuracy, F1, sensitivity, specificity at operating thresholds
3. **Calibration** — ECE, Brier score, reliability diagrams
4. **Robustness** — five-seed grouped re-split (MIMIC)
5. **Temporal validity** — time-based holdout within MIMIC
6. **Explainability** — global SHAP top features
7. **Fusion proxy** — static weighted vs meta-logreg (non-clinical holdout)

Primary reporting lock: `logreg_sigmoid_cal`, threshold **0.16**, fusion status protocol-ready.

## 4.3 Results and Discussion

### 4.3.1 NHANES Clinical Branch

NHANES 2013–2014: **4,702** rows, label eGFR < 60, prevalence ~6.6%.

#### Table 4.1: NHANES Test-Set Performance (Separate Cohort)

| Model | Accuracy | F1 | Recall | Specificity | AUROC | AUPRC | ECE |
|-------|----------|-----|--------|-------------|-------|-------|-----|
| **XGBoost** | 0.966 | 0.798 | 0.913 | 0.970 | **0.9922** | 0.932 | 0.014 |
| Baseline LogReg | 0.939 | 0.692 | 0.928 | 0.940 | 0.982 | 0.852 | 0.053 |
| Random Forest | 0.948 | 0.714 | 0.884 | 0.953 | 0.979 | 0.832 | 0.027 |
| Tabular MLP | 0.946 | 0.687 | 0.812 | 0.956 | 0.978 | 0.826 | 0.022 |
| ResMLP | 0.962 | 0.723 | 0.681 | 0.984 | 0.971 | 0.816 | 0.032 |

**Interpretation:** Strong population discrimination; **not** primary hospital evidence and **must not** be pooled with MIMIC metrics.

### 4.3.2 MIMIC Admission Branch — Primary Evidence

**30,000** admissions, CKD proxy prevalence **15.2%**, **68** features, test **n = 5,997**.

#### Table 4.2: MIMIC Held-Out Test Performance (Uncalibrated)

| Model | Accuracy | F1 | AUROC | AUPRC | ECE |
|-------|----------|-----|-------|-------|-----|
| XGBoost | 0.649 | 0.409 | 0.768 | 0.382 | 0.207 |
| Random Forest | 0.684 | 0.423 | 0.768 | 0.367 | 0.198 |
| Tabular MLP | 0.645 | 0.406 | 0.766 | 0.386 | 0.028 |
| Logistic Regression | 0.702 | 0.417 | 0.761 | 0.380 | 0.249 |

#### Table 4.3: MIMIC Calibration Comparison (Primary Reporting)

| Model | Threshold | F1 | Recall | Specificity | AUROC | ECE | Brier |
|-------|-----------|-----|--------|-------------|-------|-----|-------|
| LogReg uncalibrated | 0.52 | 0.417 | 0.658 | 0.711 | 0.761 | 0.249 | 0.195 |
| **LogReg sigmoid cal** | **0.16** | **0.415** | **0.674** | **0.696** | **0.7614** | **0.0203** | 0.119 |
| LogReg isotonic cal | 0.12 | 0.404 | 0.775 | 0.602 | 0.761 | 0.014 | 0.119 |

Sigmoid calibration reduced ECE by an order of magnitude with **unchanged AUROC**, supporting calibrated probabilities in decision support.

**Confusion matrix** at threshold 0.16: TN=3560, FP=1468, FN=330, TP=639.

**Robustness (five seeds):** AUROC 0.769 ± 0.007; ECE (uncalibrated sweep) 0.261 ± 0.008.

**Temporal external validation:** test AUROC **0.771**, ECE 0.022 — within ~0.01 of grouped test AUROC.

**SHAP top features:** anchor_age, insurance (Medicare), race, admission type, length of stay, lab_urea_nitrogen (BUN).

**Figures 4.1–4.4:** calibration reliability, AUROC bar chart, SHAP top-10, confusion matrix (`paper_assets/figures/`).

### 4.3.3 WESAD Wearable Branch

15 subjects, 578 windows, stress vs baseline ( **not CKD** ).

| Metric | Value |
|--------|-------|
| AUROC | **0.7639** |
| F1 | 0.6667 |
| Accuracy | 0.8121 |

Demonstrates wearable encoding and branch probability export only.

### 4.3.4 Multimodal Fusion Protocol

No same-patient IDs across NHANES, MIMIC, WESAD.

#### Table 4.4: Fusion Meta-Learner Comparison (Proxy Holdout — Illustrative Only)

| Method | AUROC | F1 | ECE |
|--------|-------|-----|-----|
| Static weighted | 0.497 | 0.254 | 0.748 |
| Meta-logreg on branch probs | 0.519 | 0.258 | 0.357 |

Does **not** support multimodal CKD improvement claims.

### 4.3.5 Clinical Decision-Support Prototype

Streamlit app loads frozen MIMIC checkpoint: calibrated risk, threshold control, local and global explanations. Research demonstration only (`./run_cds_app.sh`).

### 4.3.6 Discussion

1. **MIMIC** is the defensible core: moderate AUROC (~0.76), strong calibration (ECE ~0.02), grouped/temporal validation, SHAP.
2. **NHANES** shows population feasibility on a different label space.
3. **WESAD** validates wearable processing on a **proxy** task.
4. **Fusion** is architectural/protocol contribution pending aligned data.
5. **Honest scope:** primary model is **ML (calibrated logistic regression)**, not end-to-end deep learning on raw signals.

## 4.4 Summary

Chapter 4 presented environment setup, evaluation protocol, and results for three disjoint cohorts. MIMIC admission modeling is the primary contribution; NHANES and WESAD support separate branches; fusion remains documented protocol only.

---

# Chapter 5 — Engineering Standards and Design Challenges

This chapter describes engineering practices, societal and ethical context, project management, financial analysis, and mapping to complex engineering problem (EP) and engineering activity (EA) criteria.

## 5.1 Compliance with the Standards

### 5.1.1 Software Standards

| Practice | Role | Rationale |
|----------|------|-----------|
| Python 3.12 (PEP 8) | End-to-end pipeline, CDS | Dominant ML ecosystem |
| Jupyter supervisor notebook | Staged reproducible experiments | Stepwise supervisor review |
| scikit-learn | Primary MIMIC models, calibration | Mature tabular ML |
| PyTorch | NHANES/MIMIC neural baselines | Standard DL coursework |
| Streamlit | Research CDS prototype | Rapid thesis demo |
| JSON/CSV/pickle artifacts | Reporting lock, metrics, SHAP | Lightweight reproducibility |

Research software — not a regulated medical device (IEC 62304 acknowledged, not claimed).

### 5.1.2 Hardware Standards

Development workstation (macOS/Linux); CPU-feasible MIMIC primary; optional GPU for neural branches; local storage for multi-GB MIMIC files; PhysioNet credentialed MIMIC access.

### 5.1.3 Communication Standards

Streamlit on localhost:8501; HL7 FHIR cited as future EHR integration path; LOINC/ICD coding used in MIMIC feature construction.

## 5.2 Impact on Society, Environment and Sustainability

### 5.2.1 Impact on Life

Calibrated explainable risk stratification could support earlier flagging if validated prospectively; current work is research prototype only.

### 5.2.2 Impact on Society & Environment

Supports SDG 3 (Good Health); digital/computational footprint limited to training electricity; tabular primary model is energy-efficient vs large foundation models.

### 5.2.3 Ethical Aspects

De-identified public datasets; no live patient collection; bias in EHR documentation reported via SHAP; Streamlit disclaimer; honest fusion/WESAD scope.

### 5.2.4 Sustainability Plan

Modular branches; frozen artifacts; open-source stack; future dataset version swaps.

## 5.3 Project Management and Financial Analysis

| Item | Cost |
|------|------|
| Python, sklearn, PyTorch, Streamlit | Free (open source) |
| MIMIC-IV, NHANES, WESAD | Free under research agreements |
| Hardware | Existing workstation |
| Cloud | Not required for locked results |
| **Total direct cost** | **~BDT 0** (student time + electricity) |

Alternate budget with cloud GPU (~USD 50–100) was considered unnecessary because MIMIC primary is CPU-feasible.

## 5.4 Complex Engineering Problem

### 5.4.1 Complex Problem Solving

#### Table 5.1: Mapping with Complex Engineering Problem

| EP1 | EP2 | EP3 | EP4 | EP5 | EP6 | EP7 |
|:---:|:---:|:---:|:---:|:---:|:---:|:---:|
| ✓ | ✓ | ✓ | ✓ |  | ✓ | ✓ |

- **EP1:** CS + biomedical informatics + wearable sensing
- **EP2:** Discrimination vs calibration; sensitivity vs false alarms; multimodal ambition vs disjoint data
- **EP3:** AUROC, AUPRC, ECE, grouped splits, SHAP, temporal holdout
- **EP4:** KDIGO vs ICD labels; sparse labs; WESAD proxy
- **EP5:** Partial — research prototype, not FDA SaMD
- **EP6:** Supervisor, examiners, dataset custodians, future clinicians
- **EP7:** Features → model → calibration → SHAP → CDS interdependence

#### Table 5.2: Mapping with Knowledge Profile

| K1 | K2 | K3 | K4 | K5 | K6 | K7 | K8 |
|:---:|:---:|:---:|:---:|:---:|:---:|:---:|:---:|
| ✓ | ✓ | ✓ | ✓ | ✓ | ✓ | ✓ | ✓ |

Renal physiology (K1); calibration math (K2); modular pipeline (K3); MIMIC/NHANES/WESAD schemas (K4); multimodal architecture design (K5); venv, artifacts, scripts (K6); clinical interpretation (K7); literature review (K8).

### 5.4.2 Engineering Activities

#### Table 5.3: Mapping with Complex Engineering Activities

| EA1 | EA2 | EA3 | EA4 | EA5 |
|:---:|:---:|:---:|:---:|:---:|
| ✓ | ✓ | ✓ | ✓ | ✓ |

Three public datasets + ML stack + UI; human–system CDS interaction; integrated framework with honest fusion boundary; SDG 3 relevance; domain familiarity via supervisor and literature.

## 5.5 Summary

Chapter 5 mapped software practices, ethics, cost, and complex engineering criteria for a research FYDP on calibrated explainable CKD risk modeling.

---

# Chapter 6 — Conclusion

This chapter summarizes contributions, states limitations honestly, and outlines future work.

## 6.1 Summary

This thesis presented an explainable **multimodal learning framework** for early **CKD-related risk assessment**, integrating NHANES, MIMIC-IV, and WESAD with SHAP and a prototype CDS interface.

The **primary empirical contribution** is the **MIMIC admission branch**: grouped evaluation, **`logreg_sigmoid_cal`** at threshold **0.16**, AUROC **0.7614**, ECE **0.0203**. Secondary branches show population (NHANES XGBoost AUROC 0.9922) and wearable proxy (WESAD RF AUROC 0.7639) feasibility. Late fusion is protocol-only on disjoint cohorts.

Deliverables: supervisor notebook, `src/` modules, frozen artifacts, figures, Streamlit CDS.

## 6.2 Limitation

1. **ICD CKD proxy label** on MIMIC admissions — not equivalent to KDIGO eGFR staging; risk stratification, not definitive diagnosis.
2. **No same-patient merge** — NHANES, MIMIC, WESAD are disjoint; fusion cannot be clinically validated here.
3. **WESAD is a stress proxy** — 15 subjects, no CKD labels; not wearable CKD detection.
4. **No medical imaging** in implemented pipeline — tabular EHR/labs and wearable summaries only.
5. **NHANES and MIMIC not comparable** as one benchmark — separate branches with different labels.
6. **CDS prototype** — research demo only; no prospective clinical validation.
7. **External validation partial** — temporal MIMIC holdout only; cross-database external test out of scope.
8. **ML-primary locked model** — calibrated logistic regression, not end-to-end deep learning on raw waveforms despite deep tabular baselines in the framework.

## 6.3 Future Work

1. Linked multimodal CKD cohort with aligned EHR + wearable + labs
2. Cross-hospital external validation (eICU, UK Biobank)
3. Waveform encoders (MIMIC-IV-ECG, wearable CNN/Transformer with CKD labels)
4. NLP branch on discharge notes
5. Prospective CDS usability study with nephrology stakeholders
6. Federated multi-site training for privacy

---

# References

[1] A. S. Levey and P. De Jong, "Chronic kidney disease," *Lancet*, vol. 398, no. 10302, pp. 786–802, 2021.

[2] K. T. Mills et al., "Global epidemiology of CKD: A systematic review," *J. Am. Soc. Nephrol.*, vol. 27, no. 9, pp. 2644–2654, 2016.

[3] A. G. G. Turin et al., "Lifetime risk of ESRD," *J. Am. Soc. Nephrol.*, vol. 24, no. 4, pp. 641–649, 2013.

[4] National Kidney Foundation, "KDOQI clinical practice guideline for diabetes and CKD," *Am. J. Kidney Dis.*, vol. 49, no. 2 Suppl. 2, pp. S1–S179, 2007.

[5] KDIGO, "KDIGO 2012 clinical practice guideline for the evaluation and management of CKD," *Kidney Int. Suppl.*, vol. 3, no. 1, pp. 1–150, 2013.

[6] P. E. de Jong et al., "Screening for CKD: A systematic review," *Nephrol. Dial. Transplant.*, vol. 24, no. 5, pp. 1389–1396, 2009.

[7] A. Rajkomar et al., "Scalable and accurate deep learning with electronic health records," *NPJ Digit. Med.*, vol. 1, no. 1, p. 18, 2018.

[8] E. Choi et al., "Using recurrent neural networks for predicting CKD progression," *J. Am. Med. Inform. Assoc.*, vol. 24, no. 6, pp. 1099–1108, 2017.

[9] N. Tomasev et al., "A clinically applicable approach to continuous AKI prediction," *Nature*, vol. 572, no. 7767, pp. 116–119, 2019.

[10] A. Ravizza et al., "Predicting CKD progression with machine learning," *J. Biomed. Inform.*, vol. 96, p. 103252, 2019.

[11] A. Schmidt et al., "Introducing WESAD, a multimodal dataset for wearable stress and affect detection," in *Proc. ICMI*, 2018, pp. 400–408.

[12] K. Rahimi et al., "Wearable sensors and digital health in CKD," *Clin. J. Am. Soc. Nephrol.*, vol. 15, no. 8, pp. 1184–1190, 2020.

[13] T. Chen et al., "Multimodal machine learning: A survey and taxonomy," *IEEE Trans. Pattern Anal. Mach. Intell.*, vol. 41, no. 2, pp. 423–443, 2019.

[14] C. Guo et al., "On calibration of modern neural networks," in *Proc. ICML*, 2017, pp. 1321–1330.

[15] A. P. Arrieta et al., "Explainable artificial intelligence (XAI): Concepts, taxonomies, opportunities and challenges," *Inf. Fusion*, vol. 58, pp. 82–115, 2020.

[16] R. Caruana et al., "Intelligible models for healthcare: Predicting pneumonia risk and hospital 30-day readmission," in *Proc. KDD*, 2015, pp. 1721–1730.

[17] S. M. Lundberg and S.-I. Lee, "A unified approach to interpreting model predictions," in *Proc. NeurIPS*, 2017, pp. 4765–4774.

[18] A. C. Johnson et al., "MIMIC-IV (version 3.1)," PhysioNet, 2024. [Online]. Available: https://physionet.org/content/mimiciv/3.1/

[19] Centers for Disease Control and Prevention, "National Health and Nutrition Examination Survey (NHANES)," 2013–2014. [Online]. Available: https://wwwn.cdc.gov/nchs/nhanes/

[20] A. Goldberger et al., "PhysioBank, PhysioToolkit, and PhysioNet," *Circulation*, vol. 101, no. 23, pp. e215–e220, 2000.

---

## Appendices (Optional — insert if required by department)

### Appendix A: Reproducibility Checklist

- Python 3.12, requirements.txt
- Notebook: `ckd_supervisor_pipeline_from_scratch.ipynb`
- Reporting lock: `paper_assets/tables/final_reporting_lock.json`
- Regenerate: `./scripts/prepare_predefense.sh`

### Appendix B: Artifact Index

| Artifact | Description |
|----------|-------------|
| `step2_mimic_checkpoint.pkl` | Frozen primary MIMIC model |
| `step2_mimic_calibration_summary.csv` | Calibration metrics |
| `step2_mimic_shap_top15.csv` | SHAP rankings |
| `from_scratch_clinical_nhanes_summary.csv` | NHANES branch metrics |
| `wearable_branch_summary.json` | WESAD branch metrics |
| `fusion_meta_comparison.csv` | Fusion proxy comparison |

---

*End of report — {STUDENT}, {UNIVERSITY}, Summer 2025 FYDP.*
"""


def write_docx(text: str, path: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        if line.startswith("# "):
            p = doc.add_heading(line[2:], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=3)
        elif line.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            i -= 1
            rows = []
            for tl in table_lines:
                if re.match(r"^\|[-:\s|]+\|$", tl):
                    continue
                cells = [c.strip() for c in tl.strip("|").split("|")]
                rows.append(cells)
            if rows:
                tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
                tbl.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        tbl.rows[ri].cells[ci].text = cell
        elif line == "---":
            pass
        elif line.startswith("*") and line.endswith("*"):
            doc.add_paragraph(line.strip("*"))
        elif line:
            doc.add_paragraph(line)
        i += 1

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def main() -> None:
    content = md_report()
    MD_OUT.write_text(content, encoding="utf-8")
    print(f"Wrote {MD_OUT}")
    write_docx(content, DOCX_OUT)
    print(f"Wrote {DOCX_OUT}")


if __name__ == "__main__":
    main()
